"""Extracts text out of uploaded competitor documents, one chunk at a time.

A chunk is a page (PDF), a sheet (xlsx/xls), or the whole file (docx, csv,
images, and anything unsupported) - whichever unit the format already comes
divided into. Chunking rather than one atomic per-document extraction means:
- a large document doesn't need its whole text held in memory at once,
- progress can be reported while it's still running (total_chunks/processed_chunks
  on the document row),
- one bad page/sheet fails just that chunk, not the whole document - a PDF
  with 40 clean pages and 1 corrupt one still yields 40 pages of text, and the
  bad one shows up as its own error rather than sinking everything.

Each chunk independently picks "text" (read directly, format-specific library)
or "ocr" (pytesseract on a rendered page/image, when there's no text layer) -
a mixed-scan PDF can have some pages go each way.

Every format library here is optional (see requirements.txt) and imported
defensively - a missing one fails just the document types/pages it handles
with a clear "X isn't installed" error, the same way embeddings.py degrades
when sentence-transformers isn't present, rather than taking the whole module
(and everything that imports it) down at startup.
"""

from __future__ import annotations

import csv
import io
import logging
from pathlib import Path

try:
    import fitz  # PyMuPDF
except Exception:
    fitz = None
try:
    import openpyxl
except Exception:
    openpyxl = None
try:
    import xlrd
except Exception:
    xlrd = None
try:
    from docx import Document as DocxDocument
except Exception:
    DocxDocument = None
try:
    from PIL import Image
except Exception:
    Image = None
try:
    import pytesseract
except Exception:
    pytesseract = None

logger = logging.getLogger(__name__)

# A page/sheet with fewer real characters than this in its text layer is
# treated as scanned/image-only, even if PyMuPDF technically returned
# something (stray artifacts, page numbers) — not enough to trust over OCR.
MIN_NATIVE_TEXT_CHARS = 32
MAX_OCR_PAGES = 25


def _chunk(index: int, *, text: str = "", method: str | None = None, error: str | None = None) -> dict:
    return {"index": index, "text": text.strip() if text else "", "method": method, "error": error}


def total_chunks(disk_path: Path, filename: str) -> int:
    """How many chunks iter_chunks() will yield — cheap to compute upfront
    (just opens the file for its page/sheet count, no parsing of contents),
    so the document row can show real progress instead of a bare spinner."""
    suffix = Path(filename or disk_path.name).suffix.lower()
    try:
        if suffix == ".pdf" and fitz is not None:
            doc = fitz.open(str(disk_path))
            try:
                return max(1, doc.page_count)
            finally:
                doc.close()
        if suffix == ".xlsx" and openpyxl is not None:
            workbook = openpyxl.load_workbook(str(disk_path), read_only=True)
            return max(1, len(workbook.worksheets))
        if suffix == ".xls" and xlrd is not None:
            return max(1, xlrd.open_workbook(str(disk_path)).nsheets)
    except Exception:
        pass  # iter_chunks() below will surface the real error per-chunk
    return 1


def iter_chunks(disk_path: Path, filename: str):
    """Yields one {"index", "text", "method", "error"} dict per chunk, in
    order. Never raises — a handler that blows up entirely still yields a
    single failed chunk rather than losing the document with no explanation."""
    suffix = Path(filename or disk_path.name).suffix.lower()
    handler = _CHUNK_HANDLERS.get(suffix)
    if handler is None:
        yield _chunk(0, error=f"'{suffix}' has no extraction path.")
        return
    try:
        yield from handler(disk_path)
    except Exception as exc:
        logger.exception("Extraction failed for %s", filename)
        yield _chunk(0, error=str(exc))


def _pdf_chunks(disk_path: Path):
    if fitz is None:
        yield _chunk(0, error="PDF extraction needs the 'pymupdf' package, which isn't installed.")
        return

    doc = fitz.open(str(disk_path))
    try:
        for index, page in enumerate(doc):
            native = page.get_text()
            if len(native.strip()) >= MIN_NATIVE_TEXT_CHARS:
                yield _chunk(index, text=native, method="text")
                continue

            if pytesseract is None or Image is None:
                yield _chunk(
                    index,
                    error="This page has no text layer, and OCR needs the 'pytesseract' and "
                    "'Pillow' packages, which aren't installed.",
                )
                continue
            if index >= MAX_OCR_PAGES:
                yield _chunk(index, error=f"Skipped — past the {MAX_OCR_PAGES}-page OCR limit.")
                continue

            pixmap = page.get_pixmap(matrix=fitz.Matrix(2, 2))
            image = Image.open(io.BytesIO(pixmap.tobytes("png")))
            text = pytesseract.image_to_string(image)
            if text.strip():
                yield _chunk(index, text=text, method="ocr")
            else:
                yield _chunk(index, error="No text found on this page, including via OCR.")
    finally:
        doc.close()


def _xlsx_chunks(disk_path: Path):
    if openpyxl is None:
        yield _chunk(0, error="Excel extraction needs the 'openpyxl' package, which isn't installed.")
        return

    workbook = openpyxl.load_workbook(str(disk_path), read_only=True, data_only=True)
    for index, sheet in enumerate(workbook.worksheets):
        try:
            parts = [f"# {sheet.title}"]
            for row in sheet.iter_rows(values_only=True):
                cells = [str(value) for value in row if value is not None]
                if cells:
                    parts.append("\t".join(cells))
            text = "\n".join(parts)
            if text.strip():
                yield _chunk(index, text=text, method="text")
            else:
                yield _chunk(index, error=f"No data found on sheet '{sheet.title}'.")
        except Exception as exc:
            yield _chunk(index, error=str(exc))


def _xls_chunks(disk_path: Path):
    if xlrd is None:
        yield _chunk(0, error="Legacy .xls extraction needs the 'xlrd' package, which isn't installed.")
        return

    workbook = xlrd.open_workbook(str(disk_path))
    for index, sheet in enumerate(workbook.sheets()):
        try:
            parts = [f"# {sheet.name}"]
            for row_index in range(sheet.nrows):
                cells = [str(value) for value in sheet.row_values(row_index) if value not in (None, "")]
                if cells:
                    parts.append("\t".join(cells))
            text = "\n".join(parts)
            if text.strip():
                yield _chunk(index, text=text, method="text")
            else:
                yield _chunk(index, error=f"No data found on sheet '{sheet.name}'.")
        except Exception as exc:
            yield _chunk(index, error=str(exc))


def _docx_chunks(disk_path: Path):
    if DocxDocument is None:
        yield _chunk(0, error="Word extraction needs the 'python-docx' package, which isn't installed.")
        return

    document = DocxDocument(str(disk_path))
    parts = [p.text for p in document.paragraphs if p.text.strip()]
    for table in document.tables:
        for row in table.rows:
            parts.append("\t".join(cell.text for cell in row.cells))
    text = "\n".join(parts)
    if text.strip():
        yield _chunk(0, text=text, method="text")
    else:
        yield _chunk(0, error="No text found in the document.")


def _csv_chunks(disk_path: Path):
    with open(disk_path, "r", encoding="utf-8", errors="replace", newline="") as handle:
        rows = list(csv.reader(handle))
    text = "\n".join("\t".join(row) for row in rows)
    if text.strip():
        yield _chunk(0, text=text, method="text")
    else:
        yield _chunk(0, error="No data found in the file.")


def _image_chunks(disk_path: Path):
    if pytesseract is None or Image is None:
        yield _chunk(0, error="Image OCR needs the 'pytesseract' and 'Pillow' packages, which aren't installed.")
        return

    image = Image.open(disk_path)
    text = pytesseract.image_to_string(image)
    if text.strip():
        yield _chunk(0, text=text, method="ocr")
    else:
        yield _chunk(0, error="No text found in the image.")


def _doc_chunks(disk_path: Path):
    yield _chunk(0, error="Legacy .doc isn't supported yet — re-save as .docx and re-upload.")


_CHUNK_HANDLERS = {
    ".pdf": _pdf_chunks,
    ".docx": _docx_chunks,
    ".doc": _doc_chunks,
    ".xlsx": _xlsx_chunks,
    ".xls": _xls_chunks,
    ".csv": _csv_chunks,
    ".png": _image_chunks,
    ".jpg": _image_chunks,
    ".jpeg": _image_chunks,
}
